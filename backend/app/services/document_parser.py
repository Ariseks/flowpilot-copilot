from io import BytesIO

from docx import Document
from pypdf import PdfReader


TEXT_SUFFIXES = {".txt", ".md", ".csv"}


def extract_text(content: bytes, suffix: str) -> str:
    """提取受支持文件的纯文本；仅处理非加密、文本型文档。"""
    if suffix in TEXT_SUFFIXES:
        try:
            return content.decode("utf-8").strip()
        except UnicodeDecodeError as exc:
            raise ValueError("txt、md、csv 文件必须使用 UTF-8 编码") from exc
    if suffix == ".pdf":
        try:
            reader = PdfReader(BytesIO(content))
            if reader.is_encrypted:
                raise ValueError("暂不支持加密 PDF")
            text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError("PDF 解析失败，请确认文件未损坏且包含可提取文本") from exc
        if not text:
            raise ValueError("PDF 未提取到文本；扫描件或图片型 PDF 暂不支持 OCR")
        return text
    if suffix == ".docx":
        try:
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_cells = [
                cell.text.strip()
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                if cell.text.strip()
            ]
            text = "\n".join(paragraphs + table_cells).strip()
        except Exception as exc:
            raise ValueError("DOCX 解析失败，请确认文件未损坏") from exc
        if not text:
            raise ValueError("DOCX 未提取到可用文本")
        return text
    raise ValueError("不支持的文件类型")
